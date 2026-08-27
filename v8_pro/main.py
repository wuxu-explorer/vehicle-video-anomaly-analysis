# -*- coding: utf-8 -*-

"""
车辆运营与报警月度报告 V8.6
--------------------------------
功能：
1. 自动扫描 data 文件夹
2. 自动识别两个不同月份的月度报警 Excel
3. 自动读取：
   - 报警详情
   - 行驶里程
4. 自动计算：
   - 运行车辆数
   - 总运行里程
   - 总报警次数
   - 每100公里报警次数
   - 报警类型趋势
   - 各单位报警趋势
   - 典型违规车辆组 TOP5
   - 重点风险车辆组 TOP10
5. 视频异常数据可选
6. 自动识别：
   - DMS
   - ADAS
   - DSC
   - 普通摄像头
7. 自动生成：
   - Excel
   - Word
   - PNG 图表

运行：
    py -3.13 v8_pro\main.py

依赖：
    pip install pandas openpyxl numpy matplotlib python-docx
"""

from pathlib import Path
import re
import shutil
import warnings

import numpy as np
import pandas as pd

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

from openpyxl import load_workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


warnings.filterwarnings("ignore")


# ============================================================
# 1. 路径配置
# ============================================================

BASE_DIR = Path(__file__).resolve().parent

DATA_DIR = BASE_DIR / "data"
OUTPUT_DIR = BASE_DIR / "output"
CHART_DIR = OUTPUT_DIR / "charts"

EXCEL_REPORT = OUTPUT_DIR / "车辆运营与报警月度分析报告.xlsx"
WORD_REPORT = OUTPUT_DIR / "车辆运营与报警月度分析报告.docx"

DATA_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


# ============================================================
# 2. 报警类型统一名称
# ============================================================

ALARM_ALIASES = {
    "车距过近": [
        "车距过近",
        "车距太近",
        "跟车过近",
        "前车距离过近",
        "前车过近",
        "车距不足",
    ],

    "疲劳驾驶": [
        "疲劳驾驶",
        "疲劳",
        "疲劳报警",
    ],

    "驾驶员分心": [
        "驾驶员分心",
        "分心驾驶",
        "分心",
        "注意力分散",
    ],

    "无驾驶员": [
        "无驾驶员",
        "驾驶员缺失",
        "驾驶员不在位",
        "未检测到驾驶员",
    ],

    "违规打电话": [
        "违规打电话",
        "打电话",
        "使用手机",
        "电话",
        "手机",
    ],

    "违规抽烟": [
        "违规抽烟",
        "抽烟",
        "吸烟",
    ],

    "碰撞报警": [
        "碰撞报警",
        "碰撞",
        "前方碰撞",
        "碰撞预警",
    ],

    "左右盲区报警": [
        "左右盲区报警",
        "盲区报警",
        "左盲区",
        "右盲区",
        "盲区",
    ],

    "驾驶员打哈欠": [
        "驾驶员打哈欠",
        "打哈欠",
        "哈欠",
    ],
}


# ============================================================
# 3. 单位识别
# ============================================================

UNITS = {
    "采矿": [
        "采矿",
        "采矿车间",
        "采矿部",
    ],

    "选矿": [
        "选矿",
        "选矿车间",
        "选矿部",
    ],

    "运输": [
        "运输",
        "运输车间",
        "运输部",
    ],

    "维修": [
        "维修",
        "维修车间",
        "维修部",
    ],

    "工程": [
        "工程",
        "工程部",
    ],

    "其他": [
        "其他",
    ],
}


# ============================================================
# 4. 视频系统优先级
# ============================================================

VIDEO_PRIORITY = {
    "DMS": 4,
    "ADAS": 3,
    "DSC": 2,
    "普通摄像头": 1,
}


CORE_VIDEO_SYSTEMS = [
    "DMS",
    "ADAS",
    "DSC",
]


# ============================================================
# 5. 工具函数
# ============================================================

def txt(value):
    if value is None:
        return ""

    try:
        if pd.isna(value):
            return ""
    except Exception:
        pass

    return str(value).strip()


def norm(value):
    return re.sub(r"\s+", "", txt(value)).lower()


def safe_int(value):
    try:
        if pd.isna(value):
            return 0
        return int(round(float(value)))
    except Exception:
        return 0


def safe_float(value):
    try:
        if pd.isna(value):
            return 0.0
        return float(value)
    except Exception:
        return 0.0


def fmt_int(value):
    try:
        return f"{int(round(float(value))):,}"
    except Exception:
        return txt(value)


def fmt_num(value, digits=2):
    try:
        return f"{float(value):,.{digits}f}"
    except Exception:
        return txt(value)


def fmt_pct(value):
    try:
        return f"{float(value) * 100:.2f}%"
    except Exception:
        return "0.00%"


def trend(new_value, old_value):
    try:
        new_value = float(new_value)
        old_value = float(old_value)
    except Exception:
        return "无法比较"

    if abs(old_value) < 1e-12:
        if abs(new_value) < 1e-12:
            return "基本持平"
        return "上升"

    ratio = abs(new_value - old_value) / abs(old_value)

    if ratio <= 1e-6:
        return "基本持平"

    return "上升" if new_value > old_value else "下降"


def change_word(value):
    value = safe_float(value)

    if value > 0:
        return "增加"

    if value < 0:
        return "减少"

    return "基本持平"


# ============================================================
# 6. 报警类型识别
# ============================================================

def alarm_name(value):

    x = norm(value)

    for name, aliases in ALARM_ALIASES.items():

        for alias in aliases:

            if norm(alias) and norm(alias) in x:
                return name

    return txt(value)


# ============================================================
# 7. 单位识别
# ============================================================

def unit_of(value):

    x = norm(value)

    for unit, aliases in UNITS.items():

        for alias in aliases:

            if norm(alias) and norm(alias) in x:
                return unit

    return "其他"


# ============================================================
# 8. Excel列识别
# ============================================================

def find_column(df, candidates):

    normalized = {
        norm(c): c
        for c in df.columns
    }

    # 第一轮：完全匹配
    for candidate in candidates:

        key = norm(candidate)

        if key in normalized:
            return normalized[key]

    # 第二轮：包含匹配
    for candidate in candidates:

        key = norm(candidate)

        for col_norm, col in normalized.items():

            if key in col_norm or col_norm in key:
                return col

    return None


# ============================================================
# 9. 读取月度文件
# ============================================================

def read_month(path):

    path = Path(path)

    if not path.exists():
        raise FileNotFoundError(
            f"找不到文件：{path}"
        )

    excel = pd.ExcelFile(path)

    alarm_sheet = find_sheet(
        excel.sheet_names,
        [
            "报警详情",
            "报警明细",
            "报警数据",
        ]
    )

    mileage_sheet = find_sheet(
        excel.sheet_names,
        [
            "行驶里程",
            "运行里程",
            "里程",
        ]
    )

    if alarm_sheet is None:
        raise ValueError(
            f"{path.name} 缺少报警详情工作表"
        )

    if mileage_sheet is None:
        raise ValueError(
            f"{path.name} 缺少行驶里程工作表"
        )

    alarm_df = pd.read_excel(
        path,
        sheet_name=alarm_sheet
    )

    mileage_df = pd.read_excel(
        path,
        sheet_name=mileage_sheet
    )

    alarm_df.columns = [
        txt(c)
        for c in alarm_df.columns
    ]

    mileage_df.columns = [
        txt(c)
        for c in mileage_df.columns
    ]

    return alarm_df, mileage_df


def find_sheet(sheet_names, candidates):

    normalized = {
        norm(s): s
        for s in sheet_names
    }

    for candidate in candidates:

        key = norm(candidate)

        if key in normalized:
            return normalized[key]

    for candidate in candidates:

        key = norm(candidate)

        for s_norm, s in normalized.items():

            if key in s_norm or s_norm in key:
                return s

    return None


# ============================================================
# 10. 里程统计
# ============================================================

def mileage_summary(df):

    mileage_col = find_column(
        df,
        [
            "行驶里程(km)",
            "行驶里程（km）",
            "行驶里程",
            "运行里程",
            "总里程",
            "总运行里程",
        ]
    )

    if mileage_col is None:

        raise ValueError(
            "行驶里程表没有找到里程字段。\n"
            f"当前字段：{list(df.columns)}"
        )

    mileage = pd.to_numeric(
        df[mileage_col]
        .astype(str)
        .str.replace(",", "", regex=False),
        errors="coerce"
    ).fillna(0)

    vehicle_col = find_column(
        df,
        [
            "车牌号码",
            "车牌号",
            "车牌",
            "车辆编号",
            "车辆名称",
        ]
    )

    if vehicle_col:

        vehicles = int(
            df[vehicle_col]
            .map(txt)
            .ne("")
            .sum()
        )

    else:

        vehicles = len(df)

    return {
        "vehicles": vehicles,
        "mileage": float(mileage.sum()),
    }


# ============================================================
# 11. 报警统计
# ============================================================

def alarm_summary(df):

    alarm_col = find_column(
        df,
        [
            "报警类型",
            "报警名称",
            "事件类型",
            "事件名称",
            "状态名称",
        ]
    )

    if alarm_col is None:

        raise ValueError(
            "报警详情表没有找到报警类型字段。\n"
            f"当前字段：{list(df.columns)}"
        )

    names = df[alarm_col].map(alarm_name)

    counts = names.value_counts()

    total = len(df)

    rows = []

    for name in ALARM_ALIASES:

        count = int(
            counts.get(name, 0)
        )

        ratio = (
            count / total
            if total
            else 0
        )

        rows.append(
            {
                "报警类型": name,
                "报警次数": count,
                "报警占比": ratio,
            }
        )

    return pd.DataFrame(rows)


# ============================================================
# 12. 各单位报警统计
# ============================================================

def unit_summary(df):

    alarm_col = find_column(
        df,
        [
            "报警类型",
            "报警名称",
            "事件类型",
            "事件名称",
            "状态名称",
        ]
    )

    group_col = find_column(
        df,
        [
            "归属车组",
            "归属分组",
            "车辆组",
            "车组",
            "班组",
            "所属车组",
        ]
    )

    if alarm_col is None or group_col is None:

        return pd.DataFrame(
            columns=[
                "单位",
                "报警次数",
                "报警占比",
            ]
        )

    temp = df.copy()

    temp["_alarm_type"] = (
        temp[alarm_col]
        .map(alarm_name)
    )

    temp["_unit"] = (
        temp[group_col]
        .map(unit_of)
    )

    temp = temp[
        temp["_alarm_type"].isin(
            ALARM_ALIASES.keys()
        )
    ]

    if temp.empty:

        return pd.DataFrame(
            columns=[
                "单位",
                "报警次数",
                "报警占比",
            ]
        )

    result = (
        temp.groupby("_unit")
        .size()
        .reset_index(
            name="报警次数"
        )
        .rename(
            columns={
                "_unit": "单位"
            }
        )
    )

    total = len(df)

    result["报警占比"] = (
        result["报警次数"] / total
        if total
        else 0
    )

    result = result.sort_values(
        "报警次数",
        ascending=False
    )

    return result.reset_index(drop=True)


# ============================================================
# 13. 典型违规车组
# ============================================================

def top_violation_tables(df):

    alarm_col = find_column(
        df,
        [
            "报警类型",
            "报警名称",
            "事件类型",
            "事件名称",
            "状态名称",
        ]
    )

    group_col = find_column(
        df,
        [
            "归属车组",
            "归属分组",
            "车辆组",
            "车组",
            "班组",
            "所属车组",
        ]
    )

    empty = pd.DataFrame(
        columns=[
            "排名",
            "车组",
            "报警次数",
        ]
    )

    if alarm_col is None or group_col is None:

        return {
            "车距过近高频违规车组 TOP5": empty,
            "疲劳驾驶高频违规车组 TOP10": empty,
            "驾驶员分心高频违规车组 TOP10": empty,
            "碰撞报警高频违规车组 TOP5": empty,
        }

    temp = df.copy()

    temp["_alarm_type"] = (
        temp[alarm_col]
        .map(alarm_name)
    )

    temp["_group"] = (
        temp[group_col]
        .map(txt)
    )

    temp = temp[
        temp["_group"].ne("")
    ]

    def make_table(
        alarm_type,
        top_n
    ):

        x = temp[
            temp["_alarm_type"]
            == alarm_type
        ]

        if x.empty:
            return empty.copy()

        result = (
            x.groupby("_group")
            .size()
            .reset_index(
                name="报警次数"
            )
            .rename(
                columns={
                    "_group": "车组"
                }
            )
            .sort_values(
                "报警次数",
                ascending=False
            )
            .head(top_n)
            .reset_index(drop=True)
        )

        result.insert(
            0,
            "排名",
            range(
                1,
                len(result) + 1
            )
        )

        return result

    return {
        "车距过近高频违规车组 TOP5":
            make_table("车距过近", 5),

        "疲劳驾驶高频违规车组 TOP10":
            make_table("疲劳驾驶", 10),

        "驾驶员分心高频违规车组 TOP10":
            make_table("驾驶员分心", 10),

        "碰撞报警高频违规车组 TOP5":
            make_table("碰撞报警", 5),
    }


# ============================================================
# 14. 月度数据
# ============================================================

def month_data(path):

    alarm_df, mileage_df = read_month(path)

    mileage = mileage_summary(
        mileage_df
    )

    summary = alarm_summary(
        alarm_df
    )

    total_alarm = len(alarm_df)

    risk100 = (
        total_alarm /
        mileage["mileage"] *
        100
        if mileage["mileage"] > 0
        else 0
    )

    units = unit_summary(
        alarm_df
    )

    top_tables = top_violation_tables(
        alarm_df
    )

    return {
        "path": Path(path),
        "alarm_raw": alarm_df,
        "mileage_raw": mileage_df,

        "vehicles":
            mileage["vehicles"],

        "mileage":
            mileage["mileage"],

        "alarm_total":
            total_alarm,

        "risk100":
            risk100,

        "alarm_summary":
            summary,

        "unit_df":
            units,

        "top_tables":
            top_tables,
    }


# ============================================================
# 15. 月份识别
# ============================================================

def extract_month(path):

    name = Path(path).stem

    patterns = [
        r"(\d{1,2})月份",
        r"(\d{1,2})月",
        r"(\d{1,2})\s*月份",
    ]

    for pattern in patterns:

        match = re.search(
            pattern,
            name
        )

        if match:

            month = int(
                match.group(1)
            )

            if 1 <= month <= 12:
                return month

    return None


# ============================================================
# 16. 自动发现输入文件
# ============================================================

MONTH_FILES = {}
VIDEO_FILES = {}


def discover_inputs():

    global MONTH_FILES
    global VIDEO_FILES

    MONTH_FILES = {}
    VIDEO_FILES = {}

    print()
    print("=" * 70)
    print("正在扫描 data 文件夹...")
    print(f"数据目录：{DATA_DIR}")

    files = sorted(
        [
            p for p in DATA_DIR.iterdir()
            if p.suffix.lower()
            in [".xlsx", ".xls"]
            and not p.name.startswith("~$")
        ],
        key=lambda p: p.name
    )

    print(
        f"发现 Excel 文件：{len(files)} 个"
    )

    monthly_candidates = []

    for file in files:

        try:

            excel = pd.ExcelFile(file)

            print()
            print(
                f"检查文件：{file.name}"
            )

            print(
                f"  Sheet：{excel.sheet_names}"
            )

            alarm_sheet = find_sheet(
                excel.sheet_names,
                [
                    "报警详情",
                    "报警明细",
                ]
            )

            mileage_sheet = find_sheet(
                excel.sheet_names,
                [
                    "行驶里程",
                    "运行里程",
                ]
            )

            if (
                alarm_sheet is not None
                and mileage_sheet is not None
            ):

                month = extract_month(
                    file
                )

                if month is not None:

                    monthly_candidates.append(
                        (month, file)
                    )

                    print(
                        "  识别结果："
                        "月度运营报警数据 ✓"
                    )

                    print(
                        "  包含：报警详情 ✓"
                    )

                    print(
                        "  包含：行驶里程 ✓"
                    )

                continue

        except Exception as e:

            print(
                f"  检查失败：{e}"
            )

    # --------------------------------------------------------
    # 月度文件排序
    # --------------------------------------------------------

    monthly_candidates.sort(
        key=lambda x: x[0]
    )

    # 只保留每个月一个文件
    unique_months = {}

    for month, file in monthly_candidates:
        unique_months[month] = file

    if len(unique_months) < 2:

        raise RuntimeError(
            "\n至少需要两个不同月份的月度报警 Excel 文件。\n"
            "例如：\n"
            "  多宝山铜业6月份报警详情统计.xlsx\n"
            "  多宝山铜业7月份报警详情统计.xlsx"
        )

    months = sorted(
        unique_months.keys()
    )

    old_month = months[-2]
    new_month = months[-1]

    MONTH_FILES[old_month] = (
        unique_months[old_month]
    )

    MONTH_FILES[new_month] = (
        unique_months[new_month]
    )

    # --------------------------------------------------------
    # 视频异常文件
    # --------------------------------------------------------

    for file in files:

        if file in [
            MONTH_FILES[old_month],
            MONTH_FILES[new_month],
        ]:
            continue

        try:

            df = pd.read_excel(
                file,
                nrows=5
            )

            cols = [
                norm(c)
                for c in df.columns
            ]

            text = "".join(cols)

            if (
                "状态名称" in text
                and "状态类型" in text
                and "设备编号" in text
            ):

                if "A" not in VIDEO_FILES:
                    VIDEO_FILES["A"] = file
                elif "B" not in VIDEO_FILES:
                    VIDEO_FILES["B"] = file

        except Exception:
            pass

    print()
    print("=" * 70)
    print("自动识别结果")
    print("=" * 70)

    print(
        f"旧月份：{old_month}月 -> "
        f"{MONTH_FILES[old_month].name}"
    )

    print(
        f"新月份：{new_month}月 -> "
        f"{MONTH_FILES[new_month].name}"
    )

    print(
        "视频A："
        + (
            VIDEO_FILES["A"].name
            if "A" in VIDEO_FILES
            else "未提供"
        )
    )

    print(
        "视频B："
        + (
            VIDEO_FILES["B"].name
            if "B" in VIDEO_FILES
            else "未提供"
        )
    )

    print("=" * 70)
    print()

    return old_month, new_month


# ============================================================
# 17. 视频数据
# ============================================================

def detect_video_system(channel):

    text = norm(channel)

    if (
        "dms" in text
        or "驾驶员监控" in text
    ):
        return "DMS"

    if (
        "adas" in text
        or "高级驾驶辅助" in text
    ):
        return "ADAS"

    if (
        "dsc" in text
        or "驾驶状态" in text
    ):
        return "DSC"

    return "普通摄像头"


def load_video(path):

    df = pd.read_excel(path)

    df.columns = [
        txt(c)
        for c in df.columns
    ]

    group_col = find_column(
        df,
        [
            "归属车组",
            "归属分组",
            "车组",
        ]
    )

    channel_col = find_column(
        df,
        [
            "通道号",
            "通道",
            "channel",
            "Channel",
        ]
    )

    error_col = find_column(
        df,
        [
            "状态名称",
            "状态类型",
            "异常类型",
            "状态内容",
        ]
    )

    if group_col is None:
        df["归属车组"] = "未知"
        group_col = "归属车组"

    if channel_col is None:
        df["通道号"] = "未知"
        channel_col = "通道号"

    if error_col is None:
        df["状态名称"] = "视频异常"
        error_col = "状态名称"

    df["归属车组"] = (
        df[group_col]
        .map(txt)
    )

    df["通道号"] = (
        df[channel_col]
        .map(txt)
    )

    df["系统类型"] = (
        df["通道号"]
        .map(detect_video_system)
    )

    df["异常次数"] = 1

    summary = (
        df.groupby(
            [
                "归属车组",
                "通道号",
                "系统类型",
            ]
        )
        .size()
        .reset_index(
            name="异常次数"
        )
    )

    summary["系统优先级"] = (
        summary["系统类型"]
        .map(VIDEO_PRIORITY)
        .fillna(0)
    )

    summary = summary.sort_values(
        [
            "系统优先级",
            "异常次数",
        ],
        ascending=[
            False,
            False,
        ]
    )

    return df, summary


# ============================================================
# 18. 比较分析
# ============================================================

def compare(old, new):

    overview = pd.DataFrame(
        [
            [
                "总运行车辆数",
                old["vehicles"],
                new["vehicles"],
                new["vehicles"]
                - old["vehicles"],
                trend(
                    new["vehicles"],
                    old["vehicles"]
                ),
            ],

            [
                "总运行里程（KM）",
                old["mileage"],
                new["mileage"],
                new["mileage"]
                - old["mileage"],
                trend(
                    new["mileage"],
                    old["mileage"]
                ),
            ],

            [
                "总报警次数",
                old["alarm_total"],
                new["alarm_total"],
                new["alarm_total"]
                - old["alarm_total"],
                trend(
                    new["alarm_total"],
                    old["alarm_total"]
                ),
            ],

            [
                "每100公里报警次数",
                old["risk100"],
                new["risk100"],
                new["risk100"]
                - old["risk100"],
                trend(
                    new["risk100"],
                    old["risk100"]
                ),
            ],
        ],
        columns=[
            "指标",
            "旧月份",
            "新月份",
            "变化",
            "趋势",
        ]
    )

    old_summary = (
        old["alarm_summary"]
        .set_index("报警类型")
    )

    new_summary = (
        new["alarm_summary"]
        .set_index("报警类型")
    )

    rows = []

    for name in ALARM_ALIASES:

        old_count = int(
            old_summary
            .loc[name, "报警次数"]
        )

        new_count = int(
            new_summary
            .loc[name, "报警次数"]
        )

        old_ratio = float(
            old_summary
            .loc[name, "报警占比"]
        )

        new_ratio = float(
            new_summary
            .loc[name, "报警占比"]
        )

        rows.append(
            [
                name,
                old_count,
                new_count,
                new_count - old_count,
                old_ratio,
                new_ratio,
                (new_ratio - old_ratio) * 100,
                (
                    "上升"
                    if new_ratio > old_ratio
                    else "下降"
                    if new_ratio < old_ratio
                    else "基本持平"
                ),
            ]
        )

    trend_df = pd.DataFrame(
        rows,
        columns=[
            "报警类型",
            "旧月份次数",
            "新月份次数",
            "次数变化",
            "旧月份占比",
            "新月份占比",
            "占比变化（百分点）",
            "趋势",
        ]
    )

    # --------------------------------------------------------
    # 单位
    # --------------------------------------------------------

    old_unit = old["unit_df"]

    new_unit = new["unit_df"]

    all_units = sorted(
        set(old_unit["单位"])
        if not old_unit.empty
        else set()
    )

    if not new_unit.empty:

        all_units = sorted(
            set(all_units)
            | set(new_unit["单位"])
        )

    unit_rows = []

    old_map = (
        dict(
            zip(
                old_unit["单位"],
                old_unit["报警次数"]
            )
        )
        if not old_unit.empty
        else {}
    )

    new_map = (
        dict(
            zip(
                new_unit["单位"],
                new_unit["报警次数"]
            )
        )
        if not new_unit.empty
        else {}
    )

    for unit in all_units:

        old_count = safe_int(
            old_map.get(unit, 0)
        )

        new_count = safe_int(
            new_map.get(unit, 0)
        )

        ratio = (
            new_count / new["alarm_total"]
            if new["alarm_total"]
            else 0
        )

        unit_rows.append(
            [
                unit,
                old_count,
                new_count,
                new_count - old_count,
                ratio,
                trend(
                    new_count,
                    old_count
                ),
            ]
        )

    unit_df = pd.DataFrame(
        unit_rows,
        columns=[
            "单位",
            "旧月份报警次数",
            "新月份报警次数",
            "变化",
            "新月份占比",
            "趋势",
        ]
    )

    return (
        overview,
        trend_df,
        unit_df,
    )


# ============================================================
# 19. 风险 TOP10
# ============================================================

def risk_analysis(alarm_raw):

    alarm_col = find_column(
        alarm_raw,
        [
            "报警类型",
            "报警名称",
            "事件类型",
            "事件名称",
            "状态名称",
        ]
    )

    group_col = find_column(
        alarm_raw,
        [
            "归属车组",
            "归属分组",
            "车辆组",
            "车组",
            "班组",
        ]
    )

    columns = [
        "排名",
        "车组",
        "综合风险分",
        "风险等级",
        "车距过近",
        "疲劳驾驶",
        "驾驶员分心",
        "碰撞报警",
        "左右盲区报警",
        "违规打电话",
        "违规抽烟",
        "主要问题",
    ]

    if (
        alarm_col is None
        or group_col is None
    ):

        return pd.DataFrame(
            columns=columns
        )

    temp = alarm_raw.copy()

    temp["_alarm"] = (
        temp[alarm_col]
        .map(alarm_name)
    )

    temp["_group"] = (
        temp[group_col]
        .map(txt)
    )

    temp = temp[
        temp["_group"].ne("")
    ]

    if temp.empty:

        return pd.DataFrame(
            columns=columns
        )

    important = [
        "车距过近",
        "疲劳驾驶",
        "驾驶员分心",
        "碰撞报警",
        "左右盲区报警",
        "违规打电话",
        "违规抽烟",
    ]

    grouped = (
        temp.groupby(
            "_group"
        )["_alarm"]
        .value_counts()
        .unstack(
            fill_value=0
        )
    )

    for alarm in important:

        if alarm not in grouped.columns:
            grouped[alarm] = 0

    grouped = grouped[
        important
    ]

    # 权重
    weights = {
        "车距过近": 1.0,
        "疲劳驾驶": 1.5,
        "驾驶员分心": 1.5,
        "碰撞报警": 2.0,
        "左右盲区报警": 1.8,
        "违规打电话": 1.2,
        "违规抽烟": 1.0,
    }

    score = np.zeros(
        len(grouped)
    )

    for alarm in important:

        score += (
            grouped[alarm]
            .astype(float)
            * weights[alarm]
        )

    grouped["综合风险分"] = score

    grouped = grouped.sort_values(
        "综合风险分",
        ascending=False
    ).head(10)

    rows = []

    for group, row in grouped.iterrows():

        problems = []

        for alarm in important:

            if row[alarm] > 0:

                problems.append(
                    f"{alarm} {int(row[alarm])}次"
                )

        if row["综合风险分"] >= 20:
            level = "高风险"
        elif row["综合风险分"] >= 8:
            level = "中风险"
        else:
            level = "一般风险"

        rows.append(
            [
                group,
                round(
                    float(
                        row["综合风险分"]
                    ),
                    1
                ),
                level,
                *[
                    int(row[a])
                    for a in important
                ],
                "、".join(
                    problems[:3]
                ),
            ]
        )

    result = pd.DataFrame(
        rows,
        columns=[
            "车组",
            "综合风险分",
            "风险等级",
            *important,
            "主要问题",
        ]
    )

    result.insert(
        0,
        "排名",
        range(
            1,
            len(result) + 1
        )
    )

    return result


# ============================================================
# 20. 自动结论
# ============================================================

def auto_conclusions(
    old,
    new,
    risk_df
):

    conclusions = []

    alarm_diff = (
        new["alarm_total"]
        - old["alarm_total"]
    )

    conclusions.append(
        f"{NEW_MONTH}总报警次数较"
        f"{OLD_MONTH}"
        f"{change_word(alarm_diff)}"
        f"{fmt_int(abs(alarm_diff))}次。"
    )

    mileage_diff = (
        new["mileage"]
        - old["mileage"]
    )

    conclusions.append(
        f"{NEW_MONTH}总运行里程"
        f"{change_word(mileage_diff)}"
        f"{fmt_num(abs(mileage_diff))} KM。"
    )

    risk_diff = (
        new["risk100"]
        - old["risk100"]
    )

    conclusions.append(
        f"每100公里报警次数"
        f"{'上升' if risk_diff > 0 else '下降' if risk_diff < 0 else '基本持平'}"
        f"{fmt_num(abs(risk_diff))}。"
    )

    summary = new["alarm_summary"]

    if not summary.empty:

        top = (
            summary
            .sort_values(
                "报警次数",
                ascending=False
            )
            .iloc[0]
        )

        conclusions.append(
            f"{NEW_MONTH}报警次数最高的类型为"
            f"“{top['报警类型']}”，"
            f"共{fmt_int(top['报警次数'])}次，"
            f"占当月报警"
            f"{fmt_pct(top['报警占比'])}。"
        )

    if not risk_df.empty:

        high = risk_df[
            risk_df["风险等级"]
            == "高风险"
        ]

        focus = (
            high
            if not high.empty
            else risk_df
        )

        groups = "、".join(
            focus["车组"]
            .astype(str)
            .head(3)
        )

        conclusions.append(
            f"重点关注车组：{groups}。"
        )

    conclusions.append(
        "视频异常数据按照"
        "DMS > ADAS > DSC > 普通摄像头"
        "进行业务优先级排序。"
    )

    return conclusions


# ============================================================
# 21. 图表
# ============================================================

def setup_font():

    try:

        plt.rcParams[
            "font.sans-serif"
        ] = [
            "Microsoft YaHei",
            "SimHei",
            "Arial Unicode MS",
        ]

        plt.rcParams[
            "axes.unicode_minus"
        ] = False

    except Exception:
        pass


def save_bar_chart(
    title,
    categories,
    values_a,
    values_b,
    label_a,
    label_b,
    path,
):

    setup_font()

    plt.figure(
        figsize=(12, 6)
    )

    x = np.arange(
        len(categories)
    )

    width = 0.36

    bars1 = plt.bar(
        x - width / 2,
        values_a,
        width,
        label=label_a
    )

    bars2 = plt.bar(
        x + width / 2,
        values_b,
        width,
        label=label_b
    )

    plt.xticks(
        x,
        categories,
        rotation=25,
        ha="right"
    )

    plt.title(
        title,
        fontsize=14
    )

    plt.legend()

    plt.grid(
        axis="y",
        alpha=0.25
    )

    for bars in [
        bars1,
        bars2
    ]:

        for bar in bars:

            height = bar.get_height()

            if height > 0:

                plt.text(
                    bar.get_x()
                    + bar.get_width() / 2,
                    height,
                    f"{height:.0f}",
                    ha="center",
                    va="bottom",
                    fontsize=8
                )

    plt.tight_layout()

    plt.savefig(
        path,
        dpi=180,
        bbox_inches="tight"
    )

    plt.close()


def create_charts(
    old,
    new,
    trend_df,
    unit_df,
):

    shutil.rmtree(
        CHART_DIR,
        ignore_errors=True
    )

    CHART_DIR.mkdir(
        parents=True,
        exist_ok=True
    )

    charts = []

    # --------------------------------------------------------
    # 图1
    # --------------------------------------------------------

    path1 = (
        CHART_DIR
        / "01_车辆运营与报警对比.png"
    )

    categories = [
        "总里程（千KM）",
        "总报警次数",
        "每100公里报警",
    ]

    old_values = [
        old["mileage"] / 1000,
        old["alarm_total"],
        old["risk100"],
    ]

    new_values = [
        new["mileage"] / 1000,
        new["alarm_total"],
        new["risk100"],
    ]

    save_bar_chart(
        f"{OLD_MONTH}—{NEW_MONTH}"
        "车辆运营与报警对比",
        categories,
        old_values,
        new_values,
        OLD_MONTH,
        NEW_MONTH,
        path1,
    )

    charts.append(path1)

    # --------------------------------------------------------
    # 图2 报警类型
    # --------------------------------------------------------

    names = list(
        ALARM_ALIASES.keys()
    )

    temp = (
        trend_df
        .set_index("报警类型")
        .reindex(names)
        .fillna(0)
    )

    path2 = (
        CHART_DIR
        / "02_主要报警类型趋势.png"
    )

    save_bar_chart(
        f"{OLD_MONTH}—{NEW_MONTH}"
        "主要报警类型对比",
        names,
        temp[
            "旧月份次数"
        ].tolist(),
        temp[
            "新月份次数"
        ].tolist(),
        OLD_MONTH,
        NEW_MONTH,
        path2,
    )

    charts.append(path2)

    # --------------------------------------------------------
    # 图3 单位
    # --------------------------------------------------------

    path3 = (
        CHART_DIR
        / "03_各单位报警趋势.png"
    )

    if unit_df.empty:

        path3 = None

    else:

        save_bar_chart(
            f"{OLD_MONTH}—{NEW_MONTH}"
            "各单位报警趋势",
            unit_df["单位"].tolist(),
            unit_df[
                "旧月份报警次数"
            ].tolist(),
            unit_df[
                "新月份报警次数"
            ].tolist(),
            OLD_MONTH,
            NEW_MONTH,
            path3,
        )

    charts.append(path3)

    return charts


# ============================================================
# 22. Word辅助
# ============================================================

def set_cell_text(
    cell,
    text,
    bold=False,
    size=9,
):

    cell.text = txt(text)

    cell.vertical_alignment = (
        WD_CELL_VERTICAL_ALIGNMENT.CENTER
    )

    for paragraph in cell.paragraphs:

        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        for run in paragraph.runs:

            run.font.name = (
                "Microsoft YaHei"
            )

            run._element.rPr.rFonts.set(
                qn("w:eastAsia"),
                "微软雅黑"
            )

            run.font.size = Pt(size)
            run.bold = bold


def add_doc_table(
    doc,
    df,
    font_size=8.5,
):

    if df is None or df.empty:

        doc.add_paragraph(
            "暂无可用数据。"
        )

        return

    table = doc.add_table(
        rows=1,
        cols=len(df.columns)
    )

    table.style = "Table Grid"

    table.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )

    # 表头
    for i, col in enumerate(
        df.columns
    ):

        set_cell_text(
            table.rows[0].cells[i],
            col,
            bold=True,
            size=font_size,
        )

    # 数据
    for _, row in df.iterrows():

        cells = table.add_row().cells

        for i, value in enumerate(
            row
        ):

            if (
                isinstance(
                    value,
                    (float, np.floating)
                )
                and "占比" in str(
                    df.columns[i]
                )
            ):

                value = fmt_pct(
                    value
                )

            elif isinstance(
                value,
                (float, np.floating)
            ):

                value = fmt_num(
                    value
                )

            set_cell_text(
                cells[i],
                value,
                size=font_size,
            )


def add_heading(
    doc,
    text,
    level=1,
):

    p = doc.add_paragraph()

    run = p.add_run(
        text
    )

    run.bold = True
    run.font.name = (
        "Microsoft YaHei"
    )

    run._element.rPr.rFonts.set(
        qn("w:eastAsia"),
        "微软雅黑"
    )

    if level == 1:
        run.font.size = Pt(16)
    elif level == 2:
        run.font.size = Pt(13)
    else:
        run.font.size = Pt(11)

    return p


def add_picture(
    doc,
    path,
):

    if path is None:
        return

    if not Path(path).exists():
        return

    p = doc.add_paragraph()

    p.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    run = p.add_run()

    run.add_picture(
        str(path),
        width=Cm(15)
    )


# ============================================================
# 23. Word报告
# ============================================================

def word_report(
    old,
    new,
    overview,
    trend_df,
    unit_df,
    video_df,
    charts,
    risk_df,
    conclusions,
):

    doc = Document()

    section = doc.sections[0]

    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    normal = doc.styles["Normal"]

    normal.font.name = (
        "Microsoft YaHei"
    )

    normal._element.rPr.rFonts.set(
        qn("w:eastAsia"),
        "微软雅黑"
    )

    normal.font.size = Pt(10.5)

    # --------------------------------------------------------
    # 标题
    # --------------------------------------------------------

    title = doc.add_paragraph()

    title.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    run = title.add_run(
        "车辆运营与报警月度分析报告"
    )

    run.bold = True
    run.font.size = Pt(20)
    run.font.name = (
        "Microsoft YaHei"
    )

    subtitle = doc.add_paragraph()

    subtitle.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    subtitle.add_run(
        f"{OLD_MONTH} vs {NEW_MONTH}"
    ).font.size = Pt(13)

    doc.add_paragraph("")

    # --------------------------------------------------------
    # 一、车辆运营
    # --------------------------------------------------------

    add_heading(
        doc,
        "一、车辆运营分析",
        1
    )

    md = (
        new["mileage"]
        - old["mileage"]
    )

    ad = (
        new["alarm_total"]
        - old["alarm_total"]
    )

    rd = (
        new["risk100"]
        - old["risk100"]
    )

    p = doc.add_paragraph()

    p.add_run(
        f"{NEW_MONTH}共运行车辆"
        f"{fmt_int(new['vehicles'])}辆，"
        f"总运行里程"
        f"{fmt_num(new['mileage'])} KM。"
    )

    p.add_run(
        f"{OLD_MONTH}相比，"
        f"运行里程"
        f"{change_word(md)}"
        f"{fmt_num(abs(md))} KM；"
    )

    p.add_run(
        f"总报警次数"
        f"{change_word(ad)}"
        f"{fmt_int(abs(ad))}次；"
    )

    p.add_run(
        f"每100公里报警次数"
        f"{'上升' if rd > 0 else '下降' if rd < 0 else '基本持平'}"
        f"{fmt_num(abs(rd))}。"
    )

    add_picture(
        doc,
        charts[0]
    )

    add_doc_table(
        doc,
        overview,
        font_size=9
    )

    # --------------------------------------------------------
    # 二、报警趋势
    # --------------------------------------------------------

    add_heading(
        doc,
        "二、报警趋势分析",
        1
    )

    add_picture(
        doc,
        charts[1]
    )

    add_doc_table(
        doc,
        trend_df,
        font_size=8
    )

    if len(charts) >= 3:

        add_picture(
            doc,
            charts[2]
        )

    if not unit_df.empty:

        add_heading(
            doc,
            "各单位报警情况",
            2
        )

        add_doc_table(
            doc,
            unit_df,
            font_size=8.5
        )

    # --------------------------------------------------------
    # 三、典型违规
    # --------------------------------------------------------

    add_heading(
        doc,
        "三、典型违规案例",
        1
    )

    for title, df in (
        new["top_tables"]
        .items()
    ):

        add_heading(
            doc,
            title,
            2
        )

        add_doc_table(
            doc,
            df,
            font_size=9
        )

    # --------------------------------------------------------
    # 四、风险车辆组
    # --------------------------------------------------------

    add_heading(
        doc,
        "四、重点风险车辆组 TOP10",
        1
    )

    doc.add_paragraph(
        "综合风险评分用于重点关注和排查，"
        "不改变原始报警统计口径。"
    )

    add_doc_table(
        doc,
        risk_df,
        font_size=7.5
    )

    # --------------------------------------------------------
    # 五、视频异常
    # --------------------------------------------------------

    add_heading(
        doc,
        "五、重点视频异常",
        1
    )

    doc.add_paragraph(
        "视频异常业务优先级："
        "DMS > ADAS > DSC > 普通摄像头。"
    )

    if (
        video_df is None
        or video_df.empty
    ):

        doc.add_paragraph(
            "本次未提供视频异常 Excel，"
            "因此不影响月度运营和报警分析。"
        )

    else:

        add_doc_table(
            doc,
            video_df.head(20),
            font_size=8
        )

    # --------------------------------------------------------
    # 六、结论
    # --------------------------------------------------------

    add_heading(
        doc,
        "六、重点结论与整改建议",
        1
    )

    for conclusion in conclusions:

        p = doc.add_paragraph(
            style="List Bullet"
        )

        p.add_run(
            conclusion
        )

    doc.add_paragraph(
        ""
    )

    doc.add_paragraph(
        "编制：车辆运营与报警数据分析系统"
    )

    doc.add_paragraph(
        "报告自动生成"
    )

    doc.save(
        WORD_REPORT
    )

    return WORD_REPORT


# ============================================================
# 24. Excel报告
# ============================================================

def excel_report(
    old,
    new,
    overview,
    trend_df,
    unit_df,
    video_df,
    risk_df,
):

    if EXCEL_REPORT.exists():

        try:

            EXCEL_REPORT.unlink()

        except PermissionError:

            raise PermissionError(
                "请先关闭WPS/Excel中已经打开的旧报告："
                f"{EXCEL_REPORT}"
            )

    with pd.ExcelWriter(
        EXCEL_REPORT,
        engine="openpyxl"
    ) as writer:

        overview.to_excel(
            writer,
            sheet_name="一_车辆运营",
            index=False,
            startrow=2
        )

        trend_df.to_excel(
            writer,
            sheet_name="二_报警趋势",
            index=False,
            startrow=2
        )

        unit_df.to_excel(
            writer,
            sheet_name="二_报警趋势",
            index=False,
            startrow=8 + len(trend_df)
        )

        for title, df in (
            new["top_tables"]
            .items()
        ):

            # Excel sheet名称最多31字符
            sheet_name = title[:31]

            df.to_excel(
                writer,
                sheet_name=sheet_name,
                index=False
            )

        risk_df.to_excel(
            writer,
            sheet_name="重点风险车组",
            index=False,
            startrow=2
        )

        if (
            video_df is not None
            and not video_df.empty
        ):

            video_df.to_excel(
                writer,
                sheet_name="重点视频异常",
                index=False,
                startrow=2
            )

        old["alarm_summary"].to_excel(
            writer,
            sheet_name=f"{OLD_MONTH}月报警统计",
            index=False
        )

        new["alarm_summary"].to_excel(
            writer,
            sheet_name=f"{NEW_MONTH}月报警统计",
            index=False
        )

        old["mileage_raw"].to_excel(
            writer,
            sheet_name=f"{OLD_MONTH}月里程明细",
            index=False
        )

        new["mileage_raw"].to_excel(
            writer,
            sheet_name=f"{NEW_MONTH}月里程明细",
            index=False
        )

        old["alarm_raw"].to_excel(
            writer,
            sheet_name=f"{OLD_MONTH}月报警明细",
            index=False
        )

        new["alarm_raw"].to_excel(
            writer,
            sheet_name=f"{NEW_MONTH}月报警明细",
            index=False
        )

    # --------------------------------------------------------
    # Excel格式
    # --------------------------------------------------------

    wb = load_workbook(
        EXCEL_REPORT
    )

    header_fill = PatternFill(
        fill_type="solid",
        fgColor="D9EAF7"
    )

    thin = Side(
        style="thin",
        color="B7B7B7"
    )

    border = Border(
        left=thin,
        right=thin,
        top=thin,
        bottom=thin
    )

    for ws in wb.worksheets:

        ws.freeze_panes = "A2"

        for row in ws.iter_rows():

            for cell in row:

                cell.alignment = Alignment(
                    horizontal="center",
                    vertical="center",
                    wrap_text=True
                )

                cell.border = border

        if ws.max_row >= 1:

            for cell in ws[1]:

                cell.font = Font(
                    bold=True
                )

                cell.fill = header_fill

        # 自动列宽
        for col in range(
            1,
            ws.max_column + 1
        ):

            max_length = 0

            for row in range(
                1,
                min(ws.max_row, 100) + 1
            ):

                value = ws.cell(
                    row=row,
                    column=col
                ).value

                if value is None:
                    continue

                length = len(
                    str(value)
                )

                max_length = max(
                    max_length,
                    length
                )

            width = min(
                max(max_length + 2, 10),
                35
            )

            ws.column_dimensions[
                get_column_letter(col)
            ].width = width

    wb.save(
        EXCEL_REPORT
    )

    return EXCEL_REPORT


# ============================================================
# 25. 主函数
# ============================================================

def main():

    global OLD_MONTH
    global NEW_MONTH

    print()
    print("=" * 70)
    print(
        "车辆运营与报警月度报告 V8.6"
    )
    print(
        "自动识别增强稳定版"
    )
    print("=" * 70)

    # --------------------------------------------------------
    # 1. 自动发现
    # --------------------------------------------------------

    OLD_MONTH, NEW_MONTH = (
        discover_inputs()
    )

    # --------------------------------------------------------
    # 2. 读取两个月
    # --------------------------------------------------------

    print(
        f"\n正在读取{OLD_MONTH}月数据..."
    )

    old = month_data(
        MONTH_FILES[OLD_MONTH]
    )

    print(
        f"正在读取{NEW_MONTH}月数据..."
    )

    new = month_data(
        MONTH_FILES[NEW_MONTH]
    )

    # --------------------------------------------------------
    # 3. 视频
    # --------------------------------------------------------

    if "A" in VIDEO_FILES:

        try:

            video_a, video_sa = (
                load_video(
                    VIDEO_FILES["A"]
                )
            )

        except Exception as e:

            print(
                f"视频A读取失败：{e}"
            )

            video_a = pd.DataFrame()
            video_sa = pd.DataFrame()

    else:

        video_a = pd.DataFrame()
        video_sa = pd.DataFrame()

    if "B" in VIDEO_FILES:

        try:

            video_b, video_sb = (
                load_video(
                    VIDEO_FILES["B"]
                )
            )

        except Exception as e:

            print(
                f"视频B读取失败：{e}"
            )

            video_b = pd.DataFrame()
            video_sb = pd.DataFrame()

    else:

        video_b = pd.DataFrame()
        video_sb = pd.DataFrame()

    if (
        not video_sa.empty
        and not video_sb.empty
    ):

        video_df = pd.concat(
            [
                video_sa,
                video_sb,
            ],
            ignore_index=True
        )

    elif not video_sa.empty:

        video_df = video_sa.copy()

    elif not video_sb.empty:

        video_df = video_sb.copy()

    else:

        video_df = pd.DataFrame()

    # --------------------------------------------------------
    # 4. 月度比较
    # --------------------------------------------------------

    print(
        "\n正在进行月份对比..."
    )

    overview, trend_df, unit_df = (
        compare(
            old,
            new
        )
    )

    # --------------------------------------------------------
    # 5. 风险分析
    # --------------------------------------------------------

    print(
        "正在进行重点风险车组分析..."
    )

    risk_df = risk_analysis(
        new["alarm_raw"]
    )

    # --------------------------------------------------------
    # 6. 自动结论
    # --------------------------------------------------------

    conclusions = auto_conclusions(
        old,
        new,
        risk_df
    )

    # --------------------------------------------------------
    # 7. 图表
    # --------------------------------------------------------

    print(
        "正在生成图表..."
    )

    charts = create_charts(
        old,
        new,
        trend_df,
        unit_df
    )

    # --------------------------------------------------------
    # 8. Excel
    # --------------------------------------------------------

    print(
        "正在生成 Excel 报告..."
    )

    excel_path = excel_report(
        old,
        new,
        overview,
        trend_df,
        unit_df,
        video_df,
        risk_df
    )

    # --------------------------------------------------------
    # 9. Word
    # --------------------------------------------------------

    print(
        "正在生成 Word 报告..."
    )

    word_path = word_report(
        old,
        new,
        overview,
        trend_df,
        unit_df,
        video_df,
        charts,
        risk_df,
        conclusions
    )

    # --------------------------------------------------------
    # 10. 完成
    # --------------------------------------------------------

    print()
    print("=" * 70)
    print("报告生成完成！")
    print("=" * 70)

    print(
        f"{OLD_MONTH}月："
        f"{old['vehicles']}辆，"
        f"{fmt_num(old['mileage'])} KM，"
        f"{old['alarm_total']}次报警"
    )

    print(
        f"{NEW_MONTH}月："
        f"{new['vehicles']}辆，"
        f"{fmt_num(new['mileage'])} KM，"
        f"{new['alarm_total']}次报警"
    )

    print()
    print(
        "Excel："
        f"{excel_path}"
    )

    print(
        "Word："
        f"{word_path}"
    )

    print(
        "图表目录："
        f"{CHART_DIR}"
    )

    print("=" * 70)


# ============================================================
# 程序入口
# ============================================================

if __name__ == "__main__":
    main()